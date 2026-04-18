"""
DOCX Rendering Utility
======================
Converts structured agent JSON output into professional DOCX files
using python-docx.

Supported document types:
  - Resume/CV
  - Proposal
  - Compliance Document
  - Report
  - Generic (flexible sections)
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table

__all__ = [
    "render_resume",
    "render_proposal",
    "render_compliance_doc",
    "render_report",
    "render_generic",
]

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
FONT_NAME = "Calibri"
BODY_SIZE = Pt(11)
NAME_SIZE = Pt(14)
HEADING_SIZE = Pt(12)
SMALL_SIZE = Pt(9)

COLOR_PRIMARY = RGBColor(0x1F, 0x3A, 0x5F)   # dark navy
COLOR_SECONDARY = RGBColor(0x4A, 0x4A, 0x4A)  # dark grey
COLOR_ACCENT = RGBColor(0x2E, 0x75, 0xB6)     # blue accent
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_BLACK = RGBColor(0x00, 0x00, 0x00)
COLOR_LIGHT_GREY = RGBColor(0xD9, 0xD9, 0xD9)

HEADER_SHADE = "D9E2F3"  # light blue for table header rows


# ═══════════════════════════════════════════════════════════════════════════
# Common helper utilities
# ═══════════════════════════════════════════════════════════════════════════

def _set_margins(doc: Document, top: float = 1.0, bottom: float = 1.0,
                 left: float = 1.0, right: float = 1.0) -> None:
    """Set page margins (inches) for every section in the document."""
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)


def _set_run_font(run, name: str = FONT_NAME, size=BODY_SIZE,
                  bold: bool = False, italic: bool = False,
                  color: RGBColor | None = None) -> None:
    """Apply font settings to a single Run."""
    run.font.name = name
    run.font.size = size
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def _add_heading(doc: Document, text: str, level: int = 1,
                 with_border: bool = False) -> None:
    """Add a heading with consistent styling.

    level 0 → title (NAME_SIZE, primary colour)
    level 1 → section heading (HEADING_SIZE, primary colour, optional bottom border)
    level 2 → sub-heading (BODY_SIZE bold, secondary colour)
    """
    para = doc.add_paragraph()
    run = para.add_run(text.upper() if level == 1 else text)
    if level == 0:
        _set_run_font(run, size=NAME_SIZE, bold=True, color=COLOR_PRIMARY)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        _set_run_font(run, size=HEADING_SIZE, bold=True, color=COLOR_PRIMARY)
        para.space_before = Pt(12)
        para.space_after = Pt(4)
        if with_border:
            _add_bottom_border(para)
    elif level == 2:
        _set_run_font(run, size=BODY_SIZE, bold=True, color=COLOR_SECONDARY)
        para.space_before = Pt(6)
        para.space_after = Pt(2)
    else:
        _set_run_font(run, size=BODY_SIZE, bold=True)


def _add_bottom_border(paragraph) -> None:
    """Add a thin bottom border line under a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F3A5F")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_body_text(doc: Document, text: str, bold: bool = False,
                   italic: bool = False, alignment=None,
                   color: RGBColor | None = None,
                   size=BODY_SIZE) -> None:
    """Add a body-text paragraph."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    _set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    if alignment is not None:
        para.alignment = alignment
    para.paragraph_format.space_after = Pt(4)


def _add_bullet(doc: Document, text: str, level: int = 0) -> None:
    """Add a bullet-point paragraph."""
    para = doc.add_paragraph(style="List Bullet")
    para.clear()
    run = para.add_run(text)
    _set_run_font(run)
    if level > 0:
        para.paragraph_format.left_indent = Inches(0.25 * (level + 1))


def _add_table(doc: Document, headers: list[str], rows: list[list[str]],
               no_borders: bool = False) -> Table:
    """Add a formatted table with an optional shaded header row.

    Args:
        headers: column header strings.
        rows: list of row-data lists (each the same length as headers).
        no_borders: if True, remove all borders (useful for competency grids).
    Returns:
        The created Table object.
    """
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    hdr_cells = table.rows[0].cells
    for idx, h in enumerate(headers):
        hdr_cells[idx].text = ""
        run = hdr_cells[idx].paragraphs[0].add_run(h)
        _set_run_font(run, bold=True, color=COLOR_WHITE if not no_borders else COLOR_PRIMARY,
                      size=BODY_SIZE)
        if not no_borders:
            _shade_cell(hdr_cells[idx], HEADER_SHADE.replace("#", ""))

    # Data rows
    for r_idx, row_data in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            cells[c_idx].text = ""
            run = cells[c_idx].paragraphs[0].add_run(str(val))
            _set_run_font(run)

    if no_borders:
        _remove_table_borders(table)

    return table


def _shade_cell(cell, color_hex: str) -> None:
    """Shade a table cell with the given hex colour."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _remove_table_borders(table: Table) -> None:
    """Remove all borders from a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "none")
        element.set(qn("w:sz"), "0")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "auto")
        borders.append(element)
    tblPr.append(borders)


def _add_page_number(doc: Document) -> None:
    """Insert page numbers into the default footer of every section."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # PAGE field
        run = para.add_run()
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        run._r.append(fldChar1)

        run2 = para.add_run()
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = " PAGE "
        run2._r.append(instrText)

        run3 = para.add_run()
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run3._r.append(fldChar2)

        _set_run_font(run, size=SMALL_SIZE, color=COLOR_SECONDARY)
        _set_run_font(run2, size=SMALL_SIZE, color=COLOR_SECONDARY)
        _set_run_font(run3, size=SMALL_SIZE, color=COLOR_SECONDARY)


def _add_page_break(doc: Document) -> None:
    """Insert an explicit page break."""
    doc.add_page_break()


def _add_confidential_footer(doc: Document) -> None:
    """Add 'Confidential' + page number footer."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run_conf = para.add_run("CONFIDENTIAL  |  Page ")
        _set_run_font(run_conf, size=SMALL_SIZE, bold=True, color=COLOR_SECONDARY)

        # PAGE field
        run = para.add_run()
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        run._r.append(fldChar1)

        run2 = para.add_run()
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = " PAGE "
        run2._r.append(instrText)

        run3 = para.add_run()
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run3._r.append(fldChar2)

        _set_run_font(run, size=SMALL_SIZE, color=COLOR_SECONDARY)
        _set_run_font(run2, size=SMALL_SIZE, color=COLOR_SECONDARY)
        _set_run_font(run3, size=SMALL_SIZE, color=COLOR_SECONDARY)


def _title_page(doc: Document, lines: list[tuple[str, Any, bool]]) -> None:
    """Render a centred title page.

    Args:
        lines: list of (text, font_size, is_bold) tuples.
    """
    # Add some vertical space
    for _ in range(6):
        doc.add_paragraph()
    for text, size, bold in lines:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        _set_run_font(run, size=size, bold=bold, color=COLOR_PRIMARY)
    _add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════
# 1. Resume / CV
# ═══════════════════════════════════════════════════════════════════════════

def render_resume(data: dict, output_path: Path) -> Path:
    """Render a professional resume/CV from structured data.

    Expected *data* keys:
        name, email, phone, linkedin, location,
        summary,
        competencies (list[str]),
        experience (list[dict] with company, title, location, dates, bullets),
        education (list[dict] with institution, degree, dates, honors),
        certifications (list[str]),
        skills (dict with categories as keys, list[str] as values)
    """
    output_path = Path(output_path)
    doc = Document()
    _set_margins(doc)

    # ── Name ──
    _add_heading(doc, data.get("name", ""), level=0)

    # ── Contact info line ──
    contact_parts = []
    for key in ("email", "phone", "linkedin", "location"):
        val = data.get(key)
        if val:
            contact_parts.append(str(val))
    if contact_parts:
        _add_body_text(doc, "  |  ".join(contact_parts),
                       alignment=WD_ALIGN_PARAGRAPH.CENTER,
                       color=COLOR_SECONDARY, size=Pt(10))

    # ── Professional Summary ──
    summary = data.get("summary")
    if summary:
        _add_heading(doc, "Professional Summary", level=1, with_border=True)
        _add_body_text(doc, summary)

    # ── Core Competencies (3-column grid, no borders) ──
    competencies = data.get("competencies", [])
    if competencies:
        _add_heading(doc, "Core Competencies", level=1, with_border=True)
        # Pad to multiple of 3
        while len(competencies) % 3 != 0:
            competencies.append("")
        rows = [competencies[i:i + 3] for i in range(0, len(competencies), 3)]
        _add_table(doc, ["", "", ""], rows, no_borders=True)

    # ── Experience ──
    experience = data.get("experience", [])
    if experience:
        _add_heading(doc, "Professional Experience", level=1, with_border=True)
        for role in experience:
            # Company / Title line
            para = doc.add_paragraph()
            run_title = para.add_run(f"{role.get('title', '')}  —  {role.get('company', '')}")
            _set_run_font(run_title, bold=True, color=COLOR_PRIMARY)

            # Location / Dates
            meta_parts = []
            if role.get("location"):
                meta_parts.append(role["location"])
            if role.get("dates"):
                meta_parts.append(role["dates"])
            if meta_parts:
                _add_body_text(doc, "  |  ".join(meta_parts), italic=True,
                               color=COLOR_SECONDARY, size=Pt(10))

            for bullet in role.get("bullets", []):
                _add_bullet(doc, bullet)

    # ── Education ──
    education = data.get("education", [])
    if education:
        _add_heading(doc, "Education", level=1, with_border=True)
        for edu in education:
            para = doc.add_paragraph()
            run = para.add_run(
                f"{edu.get('degree', '')}  —  {edu.get('institution', '')}"
            )
            _set_run_font(run, bold=True)
            if edu.get("dates"):
                _add_body_text(doc, edu["dates"], italic=True,
                               color=COLOR_SECONDARY, size=Pt(10))
            if edu.get("honors"):
                _add_body_text(doc, edu["honors"], italic=True)

    # ── Certifications ──
    certifications = data.get("certifications", [])
    if certifications:
        _add_heading(doc, "Certifications", level=1, with_border=True)
        for cert in certifications:
            _add_bullet(doc, cert)

    # ── Skills ──
    skills = data.get("skills", {})
    if skills:
        _add_heading(doc, "Skills", level=1, with_border=True)
        for category, items in skills.items():
            para = doc.add_paragraph()
            run_cat = para.add_run(f"{category}: ")
            _set_run_font(run_cat, bold=True, color=COLOR_PRIMARY)
            run_items = para.add_run(", ".join(items))
            _set_run_font(run_items)

    _add_page_number(doc)
    doc.save(str(output_path))
    return output_path


# ═══════════════════════════════════════════════════════════════════════════
# 2. Proposal
# ═══════════════════════════════════════════════════════════════════════════

def render_proposal(data: dict, output_path: Path) -> Path:
    """Render a professional business proposal.

    Expected *data* keys:
        title, client_company, date,
        executive_summary,
        client_understanding (dict: company, industry, challenges, goals),
        proposed_solution (dict: overview, phases list[dict], tech_stack list[str]),
        scope (dict: in_scope list, out_of_scope list, assumptions list),
        timeline (list[dict]: milestone, date, deliverable),
        pricing (list[dict]: item, description, amount; total str),
        why_us (list[str]),
        terms (dict: validity, warranty, ip, confidentiality, cancellation),
        case_studies (list[dict]: title, summary, results) — optional,
        next_steps (str or list[str])
    """
    output_path = Path(output_path)
    doc = Document()
    _set_margins(doc)

    # ── Title page ──
    _title_page(doc, [
        (data.get("title", "Proposal"), Pt(28), True),
        ("", Pt(12), False),
        (f"Prepared for: {data.get('client_company', '')}", Pt(14), False),
        (f"Date: {data.get('date', '')}", Pt(12), False),
        ("", Pt(12), False),
        ("Prepared by Bit Rage Labour Systems", Pt(12), True),
    ])

    # ── Table of Contents ──
    _add_heading(doc, "Table of Contents", level=1, with_border=True)
    toc_items = [
        "1. Executive Summary",
        "2. Client Understanding",
        "3. Proposed Solution",
        "4. Scope of Work",
        "5. Timeline",
        "6. Pricing",
        "7. Why Us",
        "8. Terms & Conditions",
    ]
    if data.get("case_studies"):
        toc_items.append("9. Case Studies")
        toc_items.append("10. Next Steps")
    else:
        toc_items.append("9. Next Steps")
    for item in toc_items:
        _add_body_text(doc, item)
    _add_page_break(doc)

    # ── Executive Summary ──
    _add_heading(doc, "1. Executive Summary", level=1, with_border=True)
    _add_body_text(doc, data.get("executive_summary", ""))
    _add_page_break(doc)

    # ── Client Understanding ──
    _add_heading(doc, "2. Client Understanding", level=1, with_border=True)
    cu = data.get("client_understanding", {})
    for label, key in [("Company", "company"), ("Industry", "industry")]:
        val = cu.get(key)
        if val:
            para = doc.add_paragraph()
            r1 = para.add_run(f"{label}: ")
            _set_run_font(r1, bold=True)
            r2 = para.add_run(val)
            _set_run_font(r2)
    if cu.get("challenges"):
        _add_heading(doc, "Challenges", level=2)
        for c in cu["challenges"]:
            _add_bullet(doc, c)
    if cu.get("goals"):
        _add_heading(doc, "Goals", level=2)
        for g in cu["goals"]:
            _add_bullet(doc, g)
    _add_page_break(doc)

    # ── Proposed Solution ──
    _add_heading(doc, "3. Proposed Solution", level=1, with_border=True)
    ps = data.get("proposed_solution", {})
    if ps.get("overview"):
        _add_body_text(doc, ps["overview"])
    if ps.get("phases"):
        _add_heading(doc, "Phases", level=2)
        headers = ["Phase", "Description", "Duration"]
        rows = []
        for phase in ps["phases"]:
            rows.append([
                str(phase.get("name", phase.get("phase", ""))),
                str(phase.get("description", "")),
                str(phase.get("duration", "")),
            ])
        _add_table(doc, headers, rows)
    if ps.get("tech_stack"):
        _add_heading(doc, "Technology Stack", level=2)
        for t in ps["tech_stack"]:
            _add_bullet(doc, t)
    _add_page_break(doc)

    # ── Scope of Work ──
    _add_heading(doc, "4. Scope of Work", level=1, with_border=True)
    scope = data.get("scope", {})
    if scope.get("in_scope"):
        _add_heading(doc, "In Scope", level=2)
        for s in scope["in_scope"]:
            _add_bullet(doc, s)
    if scope.get("out_of_scope"):
        _add_heading(doc, "Out of Scope", level=2)
        for s in scope["out_of_scope"]:
            _add_bullet(doc, s)
    if scope.get("assumptions"):
        _add_heading(doc, "Assumptions", level=2)
        for s in scope["assumptions"]:
            _add_bullet(doc, s)
    _add_page_break(doc)

    # ── Timeline ──
    _add_heading(doc, "5. Timeline", level=1, with_border=True)
    timeline = data.get("timeline", [])
    if timeline:
        headers = ["Milestone", "Date", "Deliverable"]
        rows = [[str(m.get("milestone", "")), str(m.get("date", "")),
                 str(m.get("deliverable", ""))] for m in timeline]
        _add_table(doc, headers, rows)
    _add_page_break(doc)

    # ── Pricing ──
    _add_heading(doc, "6. Pricing", level=1, with_border=True)
    pricing = data.get("pricing", [])
    if pricing:
        headers = ["Item", "Description", "Amount"]
        rows = [[str(p.get("item", "")), str(p.get("description", "")),
                 str(p.get("amount", ""))] for p in pricing]
        total = data.get("total", "")
        if total:
            rows.append(["", "Total", str(total)])
        _add_table(doc, headers, rows)
    _add_page_break(doc)

    # ── Why Us ──
    _add_heading(doc, "7. Why Us", level=1, with_border=True)
    why_us = data.get("why_us", [])
    for idx, item in enumerate(why_us, 1):
        _add_body_text(doc, f"{idx}. {item}")
    _add_page_break(doc)

    # ── Terms & Conditions ──
    _add_heading(doc, "8. Terms & Conditions", level=1, with_border=True)
    terms = data.get("terms", {})
    for label, key in [
        ("Validity", "validity"),
        ("Warranty", "warranty"),
        ("Intellectual Property", "ip"),
        ("Confidentiality", "confidentiality"),
        ("Cancellation", "cancellation"),
    ]:
        val = terms.get(key)
        if val:
            para = doc.add_paragraph()
            r1 = para.add_run(f"{label}: ")
            _set_run_font(r1, bold=True)
            r2 = para.add_run(val)
            _set_run_font(r2)
            para.paragraph_format.space_after = Pt(6)

    # ── Case Studies (optional) ──
    section_num = 9
    case_studies = data.get("case_studies", [])
    if case_studies:
        _add_page_break(doc)
        _add_heading(doc, f"{section_num}. Case Studies", level=1, with_border=True)
        for cs in case_studies:
            _add_heading(doc, cs.get("title", ""), level=2)
            if cs.get("summary"):
                _add_body_text(doc, cs["summary"])
            if cs.get("results"):
                _add_heading(doc, "Results", level=2)
                if isinstance(cs["results"], list):
                    for r in cs["results"]:
                        _add_bullet(doc, r)
                else:
                    _add_body_text(doc, str(cs["results"]))
        section_num += 1

    # ── Next Steps ──
    _add_page_break(doc)
    _add_heading(doc, f"{section_num}. Next Steps", level=1, with_border=True)
    next_steps = data.get("next_steps", "")
    if isinstance(next_steps, list):
        for step in next_steps:
            _add_bullet(doc, step)
    else:
        _add_body_text(doc, str(next_steps))

    _add_page_number(doc)
    doc.save(str(output_path))
    return output_path


# ═══════════════════════════════════════════════════════════════════════════
# 3. Compliance Document
# ═══════════════════════════════════════════════════════════════════════════

def render_compliance_doc(data: dict, output_path: Path) -> Path:
    """Render a compliance / policy document.

    Expected *data* keys:
        title, company_name, effective_date, version,
        sections (list[dict]):
            number (str, e.g. '1.0'), title, body (str),
            subsections (list[dict]: number, title, body) — optional,
            policy_statement (str) — optional,
        acknowledgment (dict: text, signature_lines list[str]) — optional
    """
    output_path = Path(output_path)
    doc = Document()
    _set_margins(doc)

    # ── Title page ──
    _title_page(doc, [
        (data.get("title", "Compliance Document"), Pt(28), True),
        ("", Pt(12), False),
        (data.get("company_name", ""), Pt(16), False),
        ("", Pt(12), False),
        (f"Effective Date: {data.get('effective_date', '')}", Pt(12), False),
        (f"Version: {data.get('version', '1.0')}", Pt(12), False),
    ])

    # ── Table of Contents ──
    _add_heading(doc, "Table of Contents", level=1, with_border=True)
    for section in data.get("sections", []):
        num = section.get("number", "")
        title = section.get("title", "")
        _add_body_text(doc, f"{num}  {title}")
        for sub in section.get("subsections", []):
            _add_body_text(doc, f"    {sub.get('number', '')}  {sub.get('title', '')}")
    _add_page_break(doc)

    # ── Sections ──
    for section in data.get("sections", []):
        num = section.get("number", "")
        title = section.get("title", "")
        _add_heading(doc, f"{num}  {title}", level=1, with_border=True)

        if section.get("body"):
            _add_body_text(doc, section["body"])

        # Policy statement in distinct formatting (indented, italic, bordered)
        if section.get("policy_statement"):
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            run = para.add_run(section["policy_statement"])
            _set_run_font(run, italic=True, bold=True, color=COLOR_ACCENT)
            _add_bottom_border(para)

        # Subsections
        for sub in section.get("subsections", []):
            _add_heading(doc, f"{sub.get('number', '')}  {sub.get('title', '')}", level=2)
            if sub.get("body"):
                _add_body_text(doc, sub["body"])

    # ── Acknowledgment / Signature block ──
    ack = data.get("acknowledgment", {})
    if ack:
        _add_page_break(doc)
        _add_heading(doc, "Acknowledgment", level=1, with_border=True)
        if ack.get("text"):
            _add_body_text(doc, ack["text"])
        doc.add_paragraph()  # spacer
        for line in ack.get("signature_lines", []):
            para = doc.add_paragraph()
            para.add_run(f"{line}: ").bold = True
            run_line = para.add_run("_" * 40)
            _set_run_font(run_line)
            para.paragraph_format.space_after = Pt(16)

    _add_confidential_footer(doc)
    doc.save(str(output_path))
    return output_path


# ═══════════════════════════════════════════════════════════════════════════
# 4. Report
# ═══════════════════════════════════════════════════════════════════════════

def render_report(data: dict, output_path: Path) -> Path:
    """Render a professional business report.

    Expected *data* keys:
        title, prepared_for, date,
        executive_summary,
        sections (list[dict]):
            heading, body (str),
            table (list[dict]) — optional data table,
            bullets (list[str]) — optional,
        findings (list[str]) — optional,
        recommendations (list[str]) — optional,
        charts (list[str]) — optional placeholder labels
    """
    output_path = Path(output_path)
    doc = Document()
    _set_margins(doc)

    # ── Title page ──
    _title_page(doc, [
        (data.get("title", "Report"), Pt(28), True),
        ("", Pt(12), False),
        (f"Prepared for: {data.get('prepared_for', '')}", Pt(14), False),
        (f"Date: {data.get('date', '')}", Pt(12), False),
    ])

    # ── Executive Summary ──
    if data.get("executive_summary"):
        _add_heading(doc, "Executive Summary", level=1, with_border=True)
        _add_body_text(doc, data["executive_summary"])
        _add_page_break(doc)

    # ── Sections ──
    for section in data.get("sections", []):
        _add_heading(doc, section.get("heading", ""), level=1, with_border=True)
        if section.get("body"):
            _add_body_text(doc, section["body"])
        if section.get("bullets"):
            for b in section["bullets"]:
                _add_bullet(doc, b)
        # Data table: list of dicts → table
        tbl_data = section.get("table")
        if tbl_data and isinstance(tbl_data, list) and len(tbl_data) > 0:
            headers = list(tbl_data[0].keys())
            rows = [[str(row.get(h, "")) for h in headers] for row in tbl_data]
            _add_table(doc, headers, rows)

    # ── Key Findings ──
    findings = data.get("findings", [])
    if findings:
        _add_heading(doc, "Key Findings", level=1, with_border=True)
        for f in findings:
            _add_bullet(doc, f)

    # ── Recommendations ──
    recommendations = data.get("recommendations", [])
    if recommendations:
        _add_heading(doc, "Recommendations", level=1, with_border=True)
        for r in recommendations:
            _add_bullet(doc, r)

    # ── Charts placeholder ──
    charts = data.get("charts", [])
    if charts:
        _add_heading(doc, "Charts & Visualisations", level=1, with_border=True)
        for chart_label in charts:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(f"[Chart Placeholder: {chart_label}]")
            _set_run_font(run, italic=True, color=COLOR_SECONDARY, size=Pt(10))
            para.paragraph_format.space_after = Pt(24)

    _add_page_number(doc)
    doc.save(str(output_path))
    return output_path


# ═══════════════════════════════════════════════════════════════════════════
# 5. Generic Document
# ═══════════════════════════════════════════════════════════════════════════

def render_generic(title: str, sections: list[dict], output_path: Path) -> Path:
    """Render a generic document from a list of sections.

    Each section dict may contain:
        heading (str), content (str), bullets (list[str]),
        table (list[dict])
    """
    output_path = Path(output_path)
    doc = Document()
    _set_margins(doc)

    _add_heading(doc, title, level=0)
    doc.add_paragraph()  # spacer

    for section in sections:
        if section.get("heading"):
            _add_heading(doc, section["heading"], level=1, with_border=True)
        if section.get("content"):
            _add_body_text(doc, section["content"])
        if section.get("bullets"):
            for b in section["bullets"]:
                _add_bullet(doc, b)
        if section.get("table"):
            tbl_data = section["table"]
            if isinstance(tbl_data, list) and len(tbl_data) > 0:
                headers = list(tbl_data[0].keys())
                rows = [[str(row.get(h, "")) for h in headers] for row in tbl_data]
                _add_table(doc, headers, rows)

    _add_page_number(doc)
    doc.save(str(output_path))
    return output_path
